"""文档修改模块：docx 段落修改 + xlsx 单元格修改 + 自动备份。"""
from __future__ import annotations
import base64
import io
import json
import os
import random
import shutil
import stat
from datetime import datetime
from pathlib import Path

from config import PROJECTS_DIR
import config


def _resolve(project_name: str, filename: str) -> Path:
    path = PROJECTS_DIR / project_name / filename
    path = path.resolve()
    base = PROJECTS_DIR.resolve()
    if not str(path).startswith(str(base)):
        raise ValueError(f"路径不合法: {path}")
    return path


def _apply_default_format(doc) -> dict:
    """
    对全文应用默认格式规范（永久启用）：
      1. 所有段落左对齐
      2. 1.5 倍行距，段前段后 0 行
      3. 合并连续空白段落（最多保留1个）
      4. 清除 run 文本中残留的 ** 标记
    返回统计信息字典。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    stats = {"aligned": 0, "spacing_set": 0, "blank_removed": 0, "asterisk_cleaned": 0}

    # --- 1 & 2 & 4：遍历段落，设对齐/行距/清除 ** ---
    for para in doc.paragraphs:
        # 图片段落（含 w:drawing）和图题段落不强制左对齐，保留居中
        has_drawing = bool(para._p.findall('.//' + qn('w:drawing')))
        is_caption = para.style.name in ("Caption", "图题", "图注")
        if not has_drawing and not is_caption:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats["aligned"] += 1

        # 行距：1.5倍（240 twips × 1.5 = 360），段前段后 0
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), "360")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        stats["spacing_set"] += 1

        # 清除 ** 标记
        for run in para.runs:
            if "**" in run.text:
                run.text = run.text.replace("**", "")
                stats["asterisk_cleaned"] += 1

        # 移除段落内的 <w:br/> 换行符（AI 常用 \n\n 模拟段落间距，产生空白行）
        for br in para._p.findall(".//" + qn("w:br")):
            br.getparent().remove(br)
            stats.setdefault("br_removed", 0)
            stats["br_removed"] += 1

    # --- 3：移除所有空白段落（包括只含空格/制表符的段落）---
    # 注意：图片段落 text 为空但不应删除，需先检查是否含 w:drawing
    while True:
        paras = doc.paragraphs
        removed = False
        for para in paras:
            if not para.text.strip():
                if para._p.findall('.//' + qn('w:drawing')):
                    continue  # 图片段落，跳过
                el = para._element
                el.getparent().remove(el)
                stats["blank_removed"] += 1
                removed = True
                break  # 列表已变，重新扫描
        if not removed:
            break

    return stats


def _ensure_writable(path: Path) -> None:
    """若文件为只读，先去掉只读属性。"""
    current = os.stat(str(path)).st_mode
    if not (current & stat.S_IWRITE):
        os.chmod(str(path), current | stat.S_IWRITE)


def _backup(path: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = path.with_name(f"{path.stem}_backup_{ts}{path.suffix}")
    shutil.copy2(str(path), str(backup_path))
    return backup_path


# ---------------------------------------------------------------------------
# modify_docx_paragraph
# ---------------------------------------------------------------------------

def modify_docx_paragraph(
    project_name: str,
    filename: str,
    paragraph_index: int,
    new_text: str,
) -> str:
    try:
        from docx import Document
    except ImportError:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)

    path = _resolve(project_name, filename)
    if not path.exists():
        return json.dumps({"error": f"文件不存在: {filename}"}, ensure_ascii=False)

    try:
        with open(str(path), "rb") as _f:
            doc = Document(io.BytesIO(_f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开文档: {e}"}, ensure_ascii=False)

    paragraphs = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return json.dumps(
            {"error": f"段落索引 {paragraph_index} 超出范围 [0, {len(paragraphs) - 1}]"},
            ensure_ascii=False,
        )

    para = paragraphs[paragraph_index]
    old_text = para.text

    # 保留第一个 run 的格式，清空其他 runs
    if para.runs:
        first_run = para.runs[0]
        first_run.text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)

    _apply_default_format(doc)
    _ensure_writable(path)
    doc.save(str(path))

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "paragraph_index": paragraph_index,
            "old_value": old_text,
            "new_value": new_text,
        },
        ensure_ascii=False,
    )


# 中文字号 → 磅值对照表
_FONT_SIZE_MAP: dict[str, float] = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5,
}

# style_filter="body" 匹配的段落样式名
_BODY_STYLES = {"Normal", "正文", "Body Text", "Body Text 2", "Body Text 3",
                "Default Paragraph Style", "Text Body"}


def _is_heading(style_name: str) -> bool:
    return style_name.startswith("Heading") or style_name.startswith("标题")


def _is_body(style_name: str) -> bool:
    return style_name in _BODY_STYLES or style_name.startswith("Normal")


def _parse_pt(font_size: str) -> float | None:
    """将字号字符串解析为磅值，失败返回 None。"""
    s = font_size.strip()
    if s in _FONT_SIZE_MAP:
        return _FONT_SIZE_MAP[s]
    try:
        return float(s)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# normalize_docx_style
# ---------------------------------------------------------------------------

def normalize_docx_style(
    project_name: str,
    filename: str,
    paragraph_indices: list | None = None,
) -> str:
    """清除段落 run 级别的格式覆盖（粗体/颜色/斜体/下划线/高亮等），统一继承段落样式。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)

    path = _resolve(project_name, filename)
    if not path.exists():
        return json.dumps({"error": f"文件不存在: {filename}"}, ensure_ascii=False)

    try:
        with open(str(path), "rb") as _f:
            doc = Document(io.BytesIO(_f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开文档: {e}"}, ensure_ascii=False)

    paragraphs = doc.paragraphs
    total = len(paragraphs)
    indices = paragraph_indices if paragraph_indices else list(range(total))

    invalid = [i for i in indices if i < 0 or i >= total]
    if invalid:
        return json.dumps(
            {"error": f"段落索引超出范围 [0, {total - 1}]: {invalid}"},
            ensure_ascii=False,
        )

    # XML 标签：需要从 run 的 rPr 中移除的格式元素
    _REMOVE_TAGS = [
        qn("w:color"),      # 字体颜色
        qn("w:highlight"),  # 高亮背景色
        qn("w:strike"),     # 单删除线
        qn("w:dstrike"),    # 双删除线
        qn("w:vertAlign"),  # 上标/下标
        qn("w:shd"),        # 字符底纹
    ]

    runs_cleared = 0
    for idx in indices:
        para = paragraphs[idx]
        for run in para.runs:
            # 清除布尔格式：设为 None 表示继承段落样式
            run.bold = None
            run.italic = None
            run.underline = None

            # 清除 XML 级别的颜色/高亮等属性
            rPr = run._r.find(qn("w:rPr"))
            if rPr is not None:
                for tag in _REMOVE_TAGS:
                    for el in rPr.findall(tag):
                        rPr.remove(el)

            runs_cleared += 1

    _apply_default_format(doc)
    _ensure_writable(path)
    doc.save(str(path))

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "paragraphs_normalized": len(indices),
            "runs_cleared": runs_cleared,
        },
        ensure_ascii=False,
    )


# ---------------------------------------------------------------------------
# set_docx_font_style
# ---------------------------------------------------------------------------

def set_docx_font_style(
    project_name: str,
    filename: str,
    font_name: str | None = None,
    font_size: str | None = None,
    bold: bool | None = None,
    style_filter: str = "all",
    paragraph_indices: list | None = None,
) -> str:
    """对 docx 段落设置字体名称、字号、粗体。支持全文/标题/正文/指定段落四种范围。"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)

    if font_name is None and font_size is None and bold is None:
        return json.dumps({"error": "至少需要指定 font_name、font_size、bold 其中一项"}, ensure_ascii=False)

    pt_size: float | None = None
    if font_size is not None:
        pt_size = _parse_pt(font_size)
        if pt_size is None:
            return json.dumps({"error": f"无法识别字号: {font_size}，支持如 小四、四号、12、14"}, ensure_ascii=False)

    path = _resolve(project_name, filename)
    if not path.exists():
        return json.dumps({"error": f"文件不存在: {filename}"}, ensure_ascii=False)

    try:
        with open(str(path), "rb") as _f:
            doc = Document(io.BytesIO(_f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开文档: {e}"}, ensure_ascii=False)

    paragraphs = doc.paragraphs
    total = len(paragraphs)

    # 确定要处理的段落索引
    if paragraph_indices is not None:
        invalid = [i for i in paragraph_indices if i < 0 or i >= total]
        if invalid:
            return json.dumps({"error": f"段落索引超出范围 [0,{total-1}]: {invalid}"}, ensure_ascii=False)
        target_indices = paragraph_indices
    else:
        if style_filter == "heading":
            target_indices = [i for i, p in enumerate(paragraphs) if _is_heading(p.style.name)]
        elif style_filter == "body":
            target_indices = [i for i, p in enumerate(paragraphs) if _is_body(p.style.name)]
        else:
            target_indices = list(range(total))

    runs_modified = 0
    for idx in target_indices:
        para = paragraphs[idx]
        for run in para.runs:
            if bold is not None:
                run.bold = bold

            if pt_size is not None:
                run.font.size = Pt(pt_size)

            if font_name is not None:
                # ASCII 字体
                run.font.name = font_name
                # 中文字体（East Asian）：必须通过 XML 设置，否则中文字符不生效
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), font_name)

            runs_modified += 1

    _apply_default_format(doc)
    _ensure_writable(path)
    doc.save(str(path))

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "style_filter": style_filter if paragraph_indices is None else "指定段落",
            "paragraphs_modified": len(target_indices),
            "runs_modified": runs_modified,
        },
        ensure_ascii=False,
    )


# ---------------------------------------------------------------------------
# restructure_docx_paragraphs
# ---------------------------------------------------------------------------

# 样式配置：style_key → Word样式名 + 字体参数 + 缩进
_STRUCTURE_CONFIG = {
    "heading1": {"style": "Heading 1", "pt": 14, "bold": True,  "first_line_twips": 0},
    "heading2": {"style": "Heading 2", "pt": 12, "bold": True,  "first_line_twips": 0},
    "heading3": {"style": "Heading 3", "pt": 12, "bold": False, "first_line_twips": 0},
    "body":     {"style": "Normal",    "pt": 12, "bold": False, "first_line_twips": 480},  # 2字符×12pt×20twips
}


def restructure_docx_paragraphs(
    project_name: str,
    filename: str,
    paragraph_styles: dict,
) -> str:
    """
    批量重构段落结构。paragraph_styles 为 {段落索引: style_key} 映射，
    style_key 取值：heading1 / heading2 / heading3 / body。
    - 标题：使用 Word 标准 Heading 样式（支持生成目录），顶格无缩进
    - 正文：Normal 样式，首行缩进2字符
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)

    path = _resolve(project_name, filename)
    if not path.exists():
        return json.dumps({"error": f"文件不存在: {filename}"}, ensure_ascii=False)

    try:
        with open(str(path), "rb") as _f:
            doc = Document(io.BytesIO(_f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开文档: {e}"}, ensure_ascii=False)

    paragraphs = doc.paragraphs
    total = len(paragraphs)
    modified = 0

    for idx_raw, style_key in paragraph_styles.items():
        idx = int(idx_raw)
        if idx < 0 or idx >= total:
            continue
        cfg = _STRUCTURE_CONFIG.get(str(style_key).lower())
        if cfg is None:
            continue

        para = paragraphs[idx]

        # 1. 设置 Word 段落样式（TOC 识别依赖此处）
        try:
            para.style = doc.styles[cfg["style"]]
        except KeyError:
            pass  # 文档中无此样式时跳过，不影响其他格式

        # 2. run 级别字体：宋体 + 字号 + 粗体
        for run in para.runs:
            run.font.name = "宋体"
            run.font.size = Pt(cfg["pt"])
            run.bold = cfg["bold"]
            # 中文字体需额外设置 eastAsia，否则汉字不生效
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")

        # 3. 段落缩进
        pPr = para._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "0")
        if cfg["first_line_twips"] > 0:
            ind.set(qn("w:firstLine"), str(cfg["first_line_twips"]))
        else:
            ind.set(qn("w:firstLine"), "0")

        modified += 1

    _apply_default_format(doc)
    _ensure_writable(path)
    doc.save(str(path))

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "paragraphs_restructured": modified,
        },
        ensure_ascii=False,
    )


# ---------------------------------------------------------------------------
# insert_images_into_docx
# ---------------------------------------------------------------------------

def _select_positions_with_gap(paragraphs, num_images: int, min_chars: int = 55) -> list[int]:
    """
    从段落列表中随机选取 num_images 个插入位置（0-based），
    保证任意两个选中位置之间的段落文本总字数 >= min_chars。
    返回按索引从大到小排列的列表（供倒序插入使用）。
    """
    n = len(paragraphs)
    if n == 0:
        return []

    candidates = list(range(n))
    random.shuffle(candidates)

    selected: list[int] = []
    for pos in candidates:
        if len(selected) >= num_images:
            break
        ok = True
        for sel in selected:
            lo, hi = min(pos, sel), max(pos, sel)
            chars_between = sum(len(paragraphs[k].text) for k in range(lo, hi))
            if chars_between < min_chars:
                ok = False
                break
        if ok:
            selected.append(pos)

    return sorted(selected, reverse=True)


def insert_images_into_docx(
    project_name: str,
    source: str,
    target_filename: str,
    mode: str = "random",
    caption_prefix: str = "图",
) -> str:
    """
    从 source（子文件夹 或 含图片的 docx 文件）提取图片，
    调用视觉 AI 生成图名，随机或追加插入 target_filename，
    每张图下方附居中图题（如"图 1  示意图"）。
    图片宽度自动限制为页面行宽的 70%，且任意两张图之间至少有 55 字文本。
    """
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return json.dumps({"error": "python-docx 未安装"}, ensure_ascii=False)

    # ---- 1. 收集图片 (bytes, content_type, display_name) ----
    IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif"}
    MIME_MAP = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
        ".bmp": "image/bmp", ".webp": "image/webp", ".gif": "image/gif",
    }

    source_path = (PROJECTS_DIR / project_name / source).resolve()

    raw_images: list[tuple[bytes, str, str]] = []  # (data, mime, display_name)

    if source_path.is_dir():
        for f in sorted(source_path.iterdir()):
            if f.is_file() and f.suffix.lower() in IMAGE_EXTS:
                raw_images.append((f.read_bytes(), MIME_MAP[f.suffix.lower()], f.name))
    elif source_path.exists() and source_path.suffix.lower() == ".docx":
        try:
            with open(str(source_path), "rb") as _f:
                src_doc = Document(io.BytesIO(_f.read()))
            seen = set()
            for rId, rel in src_doc.part.rels.items():
                if "/image" in rel.reltype and rId not in seen:
                    seen.add(rId)
                    part = rel.target_part
                    ct = part.content_type  # e.g. image/jpeg
                    raw_images.append((part.blob, ct, rId))
        except Exception as e:
            return json.dumps({"error": f"无法提取 docx 中的图片: {e}"}, ensure_ascii=False)
    else:
        return json.dumps({"error": f"source 不存在或不支持: {source}"}, ensure_ascii=False)

    if not raw_images:
        return json.dumps({"error": "未在 source 中找到图片"}, ensure_ascii=False)

    # ---- 2. 调用视觉 AI 生成图名 ----
    from openai import OpenAI
    vision_client = OpenAI(api_key=config.API_KEY, base_url=config.BASE_URL)

    images_with_captions: list[tuple[bytes, str, str]] = []  # (data, mime, caption_text)
    for idx, (img_data, mime, display_name) in enumerate(raw_images):
        b64 = base64.b64encode(img_data).decode()
        data_url = f"data:{mime};base64,{b64}"
        try:
            resp = vision_client.chat.completions.create(
                model=config.IMAGE_MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": data_url}},
                    {"type": "text", "text": "请用不超过15个中文字简洁描述这张图片的主要内容，仅输出描述文字，不加标点"},
                ]}],
                temperature=0.3,
            )
            desc = (resp.choices[0].message.content or display_name).strip()
        except Exception:
            desc = display_name
        caption = f"{caption_prefix} {idx + 1}  {desc}"
        images_with_captions.append((img_data, mime, caption))

    # ---- 3. 打开目标文档 ----
    target_path = _resolve(project_name, target_filename)
    if not target_path.exists():
        return json.dumps({"error": f"目标文件不存在: {target_filename}"}, ensure_ascii=False)

    try:
        with open(str(target_path), "rb") as _f:
            doc = Document(io.BytesIO(_f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开目标文档: {e}"}, ensure_ascii=False)

    body = doc.element.body

    # ---- 4. 计算图片最大宽度（行宽 × 70%）----
    from docx.shared import Emu
    try:
        section = doc.sections[0]
        text_width_emu = section.page_width - section.left_margin - section.right_margin
        max_img_width_emu = int(text_width_emu * 0.7)
    except Exception:
        max_img_width_emu = int(Cm(10.0))  # 回退：10cm

    # ---- 5. 确定插入位置（倒序，避免索引偏移）----
    num_paras = len(doc.paragraphs)
    num_images = len(images_with_captions)

    if mode == "random" and num_paras > 0:
        positions = _select_positions_with_gap(doc.paragraphs, num_images, min_chars=55)
        # 若受间距约束导致位置不足，多余图片追加到末尾
        positions += [None] * (num_images - len(positions))
    else:
        positions = [None] * num_images  # 全部追加

    # ---- 6. 插入图片 + 图题（倒序保证随机位置稳定）----
    def _insert_before(ref_p, elem):
        """将 elem 插入到 ref_p 之前；ref_p 为 None 时追加到 sectPr 前。"""
        if ref_p is not None:
            ref_p.addprevious(elem)
        else:
            sect_pr = body.find(qn("w:sectPr"))
            if sect_pr is not None:
                sect_pr.addprevious(elem)
            else:
                body.append(elem)

    inserted = 0
    for (img_data, mime, caption), pos in zip(images_with_captions, positions):
        ref_p = doc.paragraphs[pos]._p if pos is not None else None

        # — 图片段落 —
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run()
        try:
            img_run.add_picture(io.BytesIO(img_data), width=Emu(max_img_width_emu))
        except Exception:
            img_run.add_picture(io.BytesIO(img_data))
        img_el = img_para._p
        body.remove(img_el)
        _insert_before(ref_p, img_el)

        # — 图题段落 —
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            cap_para.style = doc.styles["Caption"]
        except KeyError:
            pass  # 无 Caption 样式时保持默认
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(10.5)  # 五号
        cap_el = cap_para._p
        body.remove(cap_el)
        _insert_before(ref_p, cap_el)

        inserted += 1

    _ensure_writable(target_path)
    doc.save(str(target_path))

    return json.dumps(
        {
            "status": "success",
            "target": target_filename,
            "images_inserted": inserted,
            "mode": mode,
            "captions": [c for _, _, c in images_with_captions],
        },
        ensure_ascii=False,
    )


# ---------------------------------------------------------------------------
# modify_excel_cell
# ---------------------------------------------------------------------------

def modify_excel_cell(
    project_name: str,
    filename: str,
    sheet_name: str,
    cell_address: str,
    new_value: str,
) -> str:
    path = _resolve(project_name, filename)
    if not path.exists():
        return json.dumps({"error": f"文件不存在: {filename}"}, ensure_ascii=False)

    suffix = path.suffix.lower()
    if suffix == ".xls":
        return json.dumps(
            {"error": ".xls 格式为只读，请将文件另存为 .xlsx 后再修改"},
            ensure_ascii=False,
        )
    if suffix != ".xlsx":
        return json.dumps({"error": f"不支持的格式: {suffix}"}, ensure_ascii=False)

    try:
        import openpyxl
    except ImportError:
        return json.dumps({"error": "openpyxl 未安装"}, ensure_ascii=False)

    try:
        import io
        with open(str(path), "rb") as f:
            wb = openpyxl.load_workbook(io.BytesIO(f.read()))
    except Exception as e:
        return json.dumps({"error": f"无法打开文件: {e}"}, ensure_ascii=False)

    if sheet_name not in wb.sheetnames:
        return json.dumps(
            {"error": f"Sheet '{sheet_name}' 不存在，可用: {wb.sheetnames}"},
            ensure_ascii=False,
        )

    ws = wb[sheet_name]
    try:
        old_value = ws[cell_address].value
        ws[cell_address] = new_value
    except Exception as e:
        return json.dumps({"error": f"单元格操作失败: {e}"}, ensure_ascii=False)

    _ensure_writable(path)
    wb.save(str(path))

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "sheet": sheet_name,
            "cell": cell_address,
            "old_value": str(old_value) if old_value is not None else "",
            "new_value": new_value,
        },
        ensure_ascii=False,
    )
